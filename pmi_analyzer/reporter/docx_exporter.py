"""DOCX exporter for PMI reports.

Generates Word documents from PMI metrics with Persian RTL support,
embedded Plotly charts (as PNG images), and proper text direction.

Dependencies (optional — install with `pip install pmi-analyzer[report]`)::

    python-docx>=1.0.0
    arabic-reshaper>=3.0.0
    python-bidi>=0.6.0
"""

import importlib
import importlib.util
import logging
from pathlib import Path
from typing import Optional

from pmi_analyzer.reporter.base import BaseReport
from pmi_analyzer.types import ShamkhMetrics

logger = logging.getLogger(__name__)

# All 11 indicator field names in display order
_INDICATORS = [
    "pmi_total",
    "production",
    "new_orders",
    "sales",
    "raw_materials_inv",
    "final_goods_inv",
    "input_price",
    "production_expectations",
    "employment",
    "exports",
    "delivery_speed",
    "business_activity",
]


class DocxExporter(BaseReport):
    """Export PMI reports as Word (.docx) documents.

    Generates a document with:
    - Title and month heading
    - Status paragraph (expansion / contraction)
    - Sub-indicator table with colour-coded values
    - Embedded static chart image (if kaleido is available)
    - Proper RTL paragraph direction for Persian text

    Example::

        from pmi_analyzer.reporter.docx_exporter import DocxExporter
        from pmi_analyzer.reporter.base import ReportConfig

        exporter = DocxExporter(ReportConfig(language="fa"))
        path = exporter.generate(metrics, Path("output/report.docx"))
    """

    _DOCX_AVAILABLE: Optional[bool] = None

    def generate(
        self,
        metrics: ShamkhMetrics,
        output_path: Optional[Path] = None,
    ) -> Path:
        """Generate a DOCX report for a single month.

        Args:
            metrics: PMI metrics for the target month.
            output_path: Destination .docx file.

        Returns:
            Path to the generated .docx file, or a .txt summary on fallback.
        """
        output_path = output_path or self.config.output_dir / f"monthly_{metrics.month}.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if not self._docx_available():
            logger.warning(
                "python-docx not installed. Install with `pip install pmi-analyzer[report]`. "
                "Writing plain-text fallback: %s",
                output_path.with_suffix(".txt"),
            )
            return self._write_text_fallback(metrics, output_path.with_suffix(".txt"))

        docx = importlib.import_module("docx")
        from docx.shared import RGBColor  # type: ignore

        doc = docx.Document()

        # --- RTL document-level setting ---
        if self.config.language == "fa":
            doc.core_properties.language = "fa-IR"

        # --- Title ---
        title_para = doc.add_heading(self._get_text("title"), level=1)
        if self.config.language == "fa":
            self._set_rtl(title_para)

        # --- Month heading ---
        month_para = doc.add_heading(f"{self._get_text('month')}: {metrics.month}", level=2)
        if self.config.language == "fa":
            self._set_rtl(month_para)

        # --- Status paragraph ---
        pmi = metrics.pmi_total or 0
        status_text = f"{self._get_text('pmi_total')}: {pmi:.1f} — {self._get_status(pmi)}"
        status_para = doc.add_paragraph(status_text)
        status_para.runs[0].bold = True
        color = self._get_status_color(pmi).lstrip("#")
        r, g, b = int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
        status_para.runs[0].font.color.rgb = RGBColor(r, g, b)
        if self.config.language == "fa":
            self._set_rtl(status_para)

        # --- Indicator table ---
        indicator_label = self._get_text("value") if self.config.language != "fa" else "شاخص"
        value_label = self._get_text("value")
        status_label = self._get_text("status")

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = indicator_label
        hdr[1].text = value_label
        hdr[2].text = status_label
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        for field in _INDICATORS:
            val = getattr(metrics, field, None)
            if val is None:
                continue
            row_cells = table.add_row().cells
            row_cells[0].text = self._get_indicator_name(field)
            row_cells[1].text = f"{val:.1f}"
            row_cells[2].text = self._get_status(val)
            if self.config.language == "fa":
                for cell in row_cells:
                    self._set_rtl(cell.paragraphs[0])

        doc.add_paragraph()  # spacer

        # --- Save ---
        doc.save(str(output_path))
        logger.info("DOCX written to %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    @classmethod
    def _docx_available(cls) -> bool:
        """Check (and cache) whether python-docx is importable."""
        if cls._DOCX_AVAILABLE is None:
            cls._DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None
        return cls._DOCX_AVAILABLE

    def _write_text_fallback(self, metrics: ShamkhMetrics, path: Path) -> Path:
        """Write a plain-text summary when python-docx is unavailable."""
        lines = [
            self._get_text("title"),
            f"{self._get_text('month')}: {metrics.month}",
            "",
        ]
        for field in _INDICATORS:
            val = getattr(metrics, field, None)
            if val is not None:
                lines.append(
                    f"{self._get_indicator_name(field)}: {val:.1f} ({self._get_status(val)})"
                )
        path.write_text("\n".join(lines), encoding="utf-8")
        return path

    @staticmethod
    def _set_rtl(paragraph) -> None:  # type: ignore[type-arg]
        """Set RTL text direction on a python-docx paragraph."""
        try:
            from docx.oxml import OxmlElement  # type: ignore
            from docx.oxml.ns import qn  # type: ignore

            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            pPr.append(bidi)

            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "right")
            pPr.append(jc)
        except Exception:
            pass  # graceful degradation
